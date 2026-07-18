"""Document reader — extract text (and tables) from PDF / Word / Excel files in the workspace,
INCLUDING scanned PDFs via OCR.

Parsing real documents is library-heavy and error-prone — exactly the kind of thing to make a
vetted built-in rather than have a small model hand-roll. Text PDFs use PyMuPDF; a page with no
extractable text (a scan/image) is rasterized and run through Tesseract OCR. docx via python-docx,
xlsx via openpyxl, plain text read directly.
"""
from __future__ import annotations

import io
import os

from pydantic import BaseModel, Field

from engine.tools.base import Tool
from engine.tools.files import FileWorkspace

_MAX = 40000       # cap returned text so a huge document can't flood the context
_OCR_MIN = 20      # a page with fewer than this many extracted chars is treated as scanned → OCR


def _ocr_image(png_bytes: bytes) -> str:
    import pytesseract
    from PIL import Image
    return pytesseract.image_to_string(Image.open(io.BytesIO(png_bytes))).strip()


def _read_pdf(path: str) -> str:
    import fitz  # PyMuPDF
    doc = fitz.open(path)
    parts, ocr_pages = [], 0
    for i, page in enumerate(doc):
        txt = (page.get_text() or "").strip()
        if len(txt) < _OCR_MIN:                       # scanned/image page → OCR it
            try:
                pix = page.get_pixmap(dpi=200)
                ocr = _ocr_image(pix.tobytes("png"))
                if ocr:
                    txt, _ = ocr, ocr_pages
                    ocr_pages += 1
            except Exception as e:                    # OCR missing/broken — keep going
                txt = txt or f"[page {i+1}: no text and OCR failed: {e}]"
        parts.append(f"--- page {i+1} ---\n{txt}")
    doc.close()
    head = f"[{len(parts)} page(s), {ocr_pages} via OCR]\n" if ocr_pages else ""
    return head + "\n\n".join(parts)


def _read_docx(path: str) -> str:
    import docx
    d = docx.Document(path)
    lines = [p.text for p in d.paragraphs]
    for t in d.tables:                                # render tables as pipe rows
        for row in t.rows:
            lines.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(l for l in lines if l is not None)


def _read_xlsx(path: str) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    out = []
    for ws in wb.worksheets:
        out.append(f"# sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            if any(c is not None for c in row):
                out.append(" | ".join("" if c is None else str(c) for c in row))
    wb.close()
    return "\n".join(out)


def extract_document(path: str) -> str:
    ext = os.path.splitext(path)[1].lower().lstrip(".")
    if ext == "pdf":
        return _read_pdf(path)
    if ext == "docx":
        return _read_docx(path)
    if ext in ("xlsx", "xlsm"):
        return _read_xlsx(path)
    if ext in ("txt", "md", "csv", "json", "log", "tsv", "yaml", "yml", ""):
        with open(path, encoding="utf-8", errors="replace") as fh:
            return fh.read()
    raise ValueError(f"unsupported document type: .{ext}")


class ReadDocumentTool(Tool):
    name = "read_document"
    description = (
        "Extract the text (and tables) from a document in the workspace — PDF, Word (.docx), or "
        "Excel (.xlsx). Handles SCANNED PDFs automatically via OCR. Use this to read an uploaded "
        "file so you can summarize, answer questions about it, or add it to your knowledge base. "
        "Arg: name (the file's name in the workspace)."
    )

    class Params(BaseModel):
        name: str = Field(..., description="the document's file name in the workspace")

    def __init__(self, ws: FileWorkspace):
        self.ws = ws

    async def run(self, args: "ReadDocumentTool.Params") -> str:
        import asyncio
        path = self.ws.path_if_exists(args.name)
        if not path:
            names = ", ".join(f["name"] for f in self.ws.list()) or "(empty)"
            return f"read_document: no file '{args.name}'. Files in workspace: {names}."
        try:
            text = await asyncio.to_thread(extract_document, path)   # OCR can be slow — off-thread
        except Exception as e:
            return f"read_document: could not read '{args.name}' ({type(e).__name__}: {e})."
        text = (text or "").strip()
        if not text:
            return f"read_document: '{args.name}' had no extractable text (even after OCR)."
        if len(text) > _MAX:
            text = text[:_MAX] + f"\n… (truncated; extracted {len(text)} chars total)"
        return text
